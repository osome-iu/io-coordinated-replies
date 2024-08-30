import glob
import os
import matplotlib.pyplot as plt
import pandas as pd

def change_extension_of_file(path,
                             search_extension, 
                             replace_extension
                            ):
    '''
    Change the extension of files
    
    :param path: the path of file where it is located
    :param search_extension: extension to be searched
    :param replace_extension: extension to be replaced
    '''
    
    print('\n\n ----- Starting the change in extension ---- \n\n')
    
    new_path = os.path.join(path, f'*{search_extension}')

    for image in glob.glob(new_path):
        parts = image.split(os.sep)[-1]
        name = parts.split('.')[0]
        new_name = name + '.' +  replace_extension
        new_name_full = os.path.join(path, new_name)
    
        os.rename(image, new_name_full)
        
    print('\n\n ----- Ending the change in extension ---- \n\n')
    
    
def show_image(image):
    '''
    Shows image file
    
    :param image: Shows image
    '''
    
    im = plt.imread(image)
    
    plt.imshow(im)
    plt.show()
    
    
def write_to_file_row_each_line(path, file_name=None, 
                                rows=None):
    '''
    Writes each row as one line in file
    :param path: path where file is to be saved
    :param file_name: name of new file
    :param rows: list of row to be written
    '''
    if file_name != None:
        path_with_name_of_file = os.path.join(path, file_name)
    else:
        path_with_name_of_file = path

    with open(f'{path_with_name_of_file}', 'w') as f:
        for line in rows:
            f.write(f"{line}\n")
            
            
def create_folder(output_path, folder_name):
    '''
    Creates folder in output path if folder does not exists
    
    :param output_path: path where folder existence to be
    checked
    :param folder_name: name of folder
    '''
    
    path = os.path.join(output_path, folder_name)
    
    isExist = os.path.exists(path)

    if not isExist:
        os.makedirs(path)
        
    return path
        
def read_file(file):
    '''
    This function read files and return the content
    :param file: file name with location
    '''
    with open(f'{file}') as f:
        lines = f.read().splitlines()

        return lines
    
    
    
def remove_row_in_file(path, org_list, save_path=None):
    '''
    Removes unwanted rows from file
    :param path: Path where files are
    :param save_path: Path where the files to be saved
    :param org_list: List of original data without unwanted rows
    '''
    for file in glob.glob(path):
        filename = file.split(os.sep)[-1]
        ids = set(read_file(file))
        org_list = set(org_list)
        
        remaining = ids - org_list #unwanted rows
        
        if len(remaining) == 0:
            print(0)
            continue
            
        ids = ids - remaining
        temp_file = open(f'{file}', 'r+')
        
        temp_file.truncate(0)
        
        if save_path == None:
            save_path = path
            filename = None
            
        write_to_file_row_each_line(save_path, 
                                            filename, 
                                            ids)

def empty_the_file(filename):
    '''
    Empty the text file
    :param filename: name of file with path
    '''
    temp_file = open(f'{filename}', 'r+')
        
    temp_file.truncate(0)

        
            
def split_into_files(input_file, 
                     save_path,
                     split_threshold=50,
                     prefix_for_file='job_control'
                    ):
    '''
    Splits the rows of file into multiple files
    :param input_file: text file which has data
    :param split_threshold: threshold to split the rows by
    :param save_path: path where files are to be saved
    :param prefix_for_file: prefix for new file names
    '''
    
    rows = file_hp.read_file(input_file)
    index = 1
    for i in range(0, len(rows), split_threshold):
        ids_split = ids[i:i+split_threshold]
        last = i + split_threshold

        #first is the index of job
        #second and third are the index of rows
        filename = f'{prefix_for_file}_{index}_{i}_{last}.txt'

        file_hp.write_to_file_row_each_line(save_path,
                                            filename,
                                            ids_split
                                           )

        index = index + 1
        
        
        
def convert_to_docx(df, 
                    columns, 
                    filename):
    '''
    Converts dataframe to docx
    :param df: Dataframe to save
    :param columns: Name of columns to use in dataframe
    :param filename: Name of file to sav
    
    :return None
    '''
    
    from docx import Document

    doc = Document()

    df = df[columns]
    for index, row in df.iterrows():
        all_text = ''
        for index, column in enumerate(columns):
            text = row[column]
            if index == 0:
                all_text = text
                continue
                
            all_text = all_text + '***' + text
            
        doc.add_paragraph(f'[{all_text}]')

    doc.save(filename)
    
    
def convert_docx_to_csv(filename,
                        columns,
                        save_filename,
                        to='csv'
                      ):
    '''
    Converts the docx file into dataframe. 
    This function supports only two column to save
    :param filename : Docx filename to load
    :param columns: Name of columns for dataframe
    :param save_filename: Name of the file to save as
    :param to: Convert to the file form
    
    :return None
    '''
    from docx import Document

    doc = Document(filename)
    open_brac = '['
    closing_brac = ']'

    all_element = []
    print('Total data : ', len(doc.paragraphs))
    for p in doc.paragraphs:
        if p.text.strip() != "":
            element_list = (p.text).split('***')

            if len(element_list) == 1:
                continue

            element_list[0] = element_list[0].replace(open_brac, "")
            element_list[0] = element_list[0].replace(' ', "")
            element_list[1] = element_list[1].replace(closing_brac, "")

            if len(element_list) > 2:
                list_string = ' '.join(element_list[1:])
                element_list[1]  = list_string.replace('*', "")

            all_element.append([element_list[0], element_list[1]])


    df = pd.DataFrame(all_element, columns=columns)
    
    if to == 'csv':
        df.to_csv(save_filename,
                  index=False
                 )
    else:
        df.to_pickle(save_filename)
        
    print('Total data after dataframe formation: ', len(df))
    
    
def convert_to_docx(df, 
                    columns, 
                    filename):
    '''
    Converts dataframe to docx
    :param df: Dataframe to save
    :param columns: Name of columns to use in dataframe
    :param filename: Name of file to sav
    
    :return None
    '''
    
    from docx import Document

    doc = Document()

    df = df[columns]
    for index, row in df.iterrows():
        all_text = ''
        for column in columns:
            text = row[column]
            all_text = all_text + '***' + text
            
        doc.add_paragraph(f'[{all_text}]')

    doc.save(filename)
    
    
    
def convert_docx_to_csv(filename,
                       columns,
                       save_filename,
                        to='csv'
                      ):
    '''
    Converts the docx file into dataframe. 
    This function supports only two column to save
    :param filename : Docx filename to load
    :param columns: Name of columns for dataframe
    :param save_filename: Name of the file to save as
    
    :return None
    '''
    from docx import Document

    doc = Document(filename)
    open_brac = '['
    closing_brac = ']'

    all_element = []
    print('Total data : ', len(doc.paragraphs))
    for p in doc.paragraphs:
        if p.text.strip() != "":
            element_list = (p.text).split('***')

            if len(element_list) == 1:
                continue

            element_list[0] = element_list[0].replace(open_brac, "")
            element_list[0] = element_list[0].replace(' ', "")
            element_list[1] = element_list[1].replace(closing_brac, "")

            if len(element_list) > 2:
                list_string = ' '.join(element_list[1:])
                element_list[1]  = list_string.replace('*', "")

            all_element.append([element_list[0], element_list[1]])


    df = pd.DataFrame(all_element, columns=columns)
    
    if to == 'csv':
        df.to_csv(save_filename,
                  index=False
                 )
    else:
        df.to_pickle(save_filename)
        
    print('Total data after dataframe formation: ', len(df))